import os
from dotenv import load_dotenv
from openai import OpenAI
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import tiktoken
from datetime import datetime

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def criar_cabecalho_documento(doc, info_assembleia):
    """Cria o cabeçalho padrão do documento"""
    # Cabeçalho da administradora
    cabecalho = doc.add_paragraph()
    cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = cabecalho.add_run("Rua Batista de Oliveira 1164/206 - Centro - Juiz de Fora (MG)\n")
    run.font.size = Pt(10)
    
    run = cabecalho.add_run("CEP 36010-532 - PABX (32) 3228-8800\n")
    run.font.size = Pt(10)
    
    run = cabecalho.add_run("Horário de atendimento: 9h às 11h30 e 13h30 às 17h\n")
    run.font.size = Pt(10)
    
    run = cabecalho.add_run("www.contatojf.com.br")
    run.font.size = Pt(10)
    
    # Linha separadora
    doc.add_paragraph("_" * 80)
    
    # Nome do condomínio (do usuário)
    nome_condominio = doc.add_paragraph()
    nome_condominio.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = nome_condominio.add_run(info_assembleia.get('nome_condominio', 'CONDOMÍNIO'))
    run.font.size = Pt(14)
    run.font.bold = True

def criar_titulo_ata(doc, data_assembleia, tipo_assembleia="EXTRAORDINÁRIA"):
    """Cria o título da ata com data"""
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(f"ATA DA ASSEMBLEIA GERAL {tipo_assembleia} REALIZADA EM {data_assembleia}")
    run.font.size = Pt(12)
    run.font.bold = True

def extrair_info_assembleia(transcricao):
    """Extrai informações específicas da assembleia usando IA"""
    prompt = f"""
Analise esta transcrição de assembleia de condomínio e extraia as seguintes informações específicas em formato JSON:

1. data_assembleia: Data da assembleia (formato DD/MM/AAAA)
2. horario_inicio: Horário de início
3. tipo_assembleia: Se é "ORDINÁRIA" ou "EXTRAORDINÁRIA"
4. presidente_nome: Nome do presidente da mesa
5. presidente_apartamento: Apartamento do presidente
6. secretario_nome: Nome do secretário
7. secretario_apartamento: Apartamento do secretário
8. numero_presentes: Quantos presentes na assembleia
9. pautas: Lista das pautas principais discutidas
10. decisoes: Lista das principais decisões tomadas
11. votacao_resultado: Resultado das votações (favoráveis, contrários, abstenções)

Transcrição:
{transcricao[:3000]}...

Responda APENAS com um JSON válido, sem explicações adicionais.
"""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Você extrai informações específicas de transcrições de assembleias e responde apenas com JSON válido."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=1000,
        )
        
        import json
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        # Fallback com dados padrão
        return {
            "data_assembleia": datetime.now().strftime("%d/%m/%Y"),
            "horario_inicio": "19h40",
            "tipo_assembleia": "EXTRAORDINÁRIA",
            "presidente_nome": "A ser definido",
            "presidente_apartamento": "N/A",
            "secretario_nome": "A ser definido", 
            "secretario_apartamento": "N/A",
            "numero_presentes": "N/A",
            "pautas": ["Assuntos diversos"],
            "decisoes": ["Decisões a serem definidas"],
            "votacao_resultado": {"favoráveis": 0, "contrários": 0, "abstenções": 0}
        }

def criar_paragrafo_abertura(doc, info):
    """Cria o parágrafo de abertura padrão"""
    data_completa = converter_data_por_extenso(info['data_assembleia'])
    
    abertura = doc.add_paragraph()
    abertura.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Extrai pautas da transcrição se não fornecidas
    pautas_texto = "assuntos diversos"  # Valor padrão
    if 'pautas' in info and info['pautas']:
        if isinstance(info['pautas'], list):
            pautas_texto = ', '.join(info['pautas'])
        else:
            pautas_texto = str(info['pautas'])
    
    # Formata presidente (com ou sem apartamento)
    presidente_info = info['presidente_nome']
    if info.get('presidente_apartamento') and info['presidente_apartamento'] != 'N/A':
        presidente_info += f", apto. {info['presidente_apartamento']}"
    
    # Formata secretário (com ou sem apartamento)
    secretario_info = info['secretario_nome']
    if info.get('secretario_apartamento') and info['secretario_apartamento'] != 'N/A':
        secretario_info += f", apto. {info['secretario_apartamento']}"
    
    texto_abertura = f"""Aos {data_completa}, às {info['horario_inicio']}, iniciou-se a Assembleia geral {info['tipo_assembleia']} do {info.get('nome_condominio', 'Condomínio')}, CNPJ n° 47.688.457/0001-02, situado à Avenida Engenheiro Valdir Pedro Monachesi, n° 1.400, no bairro Aeroporto, em Juiz de Fora, Minas Gerais, {info.get('local_realizacao', '')}, conforme {info['numero_presentes']} presentes, para deliberar sobre a seguinte pauta: {pautas_texto}. Para presidir a assembleia foi convidado(a) {presidente_info}, e para secretariá-la, {secretario_info}. A leitura do edital de convocação foi realizada, porém, dispensada a leitura da ata da última assembleia, que teve a aprovação de todos."""
    
    abertura.add_run(texto_abertura)

def converter_data_por_extenso(data_str):
    """Converte data DD/MM/AAAA para formato por extenso"""
    try:
        dia, mes, ano = data_str.split('/')
        meses = {
            '01': 'janeiro', '02': 'fevereiro', '03': 'março', '04': 'abril',
            '05': 'maio', '06': 'junho', '07': 'julho', '08': 'agosto',
            '09': 'setembro', '10': 'outubro', '11': 'novembro', '12': 'dezembro'
        }
        
        return f"{dia} dias do mês de {meses[mes]} de {ano}"
    except:
        return "data a ser definida"

def dividir_texto_em_blocos(texto, max_tokens=3500):
    """Divide texto em blocos menores"""
    tokenizer = tiktoken.get_encoding("cl100k_base")
    tokens = tokenizer.encode(texto)
    blocos = []
    for i in range(0, len(tokens), max_tokens):
        bloco_tokens = tokens[i:i+max_tokens]
        bloco_texto = tokenizer.decode(bloco_tokens)
        blocos.append(bloco_texto)
    return blocos

def gerar_conteudo_formal(bloco_texto, info_assembleia):
    """Gera conteúdo formal baseado no modelo padrão"""
    prompt = f"""
Você é um redator profissional de atas de assembleia de condomínio seguindo o padrão da Contato Administração.

IMPORTANTE: Transforme este trecho da transcrição em texto formal seguindo EXATAMENTE o estilo da ata modelo fornecida:

CARACTERÍSTICAS DO MODELO:
- Texto narrativo em terceira pessoa
- Linguagem formal e técnica 
- Parágrafos longos e detalhados
- Sempre mencionar valores monetários por extenso entre parênteses
- Incluir detalhes técnicos como CNPJ, endereços, frações ideais
- Usar títulos em negrito para seções principais
- Manter sequência cronológica dos fatos

CONTEXTO DA ASSEMBLEIA:
- Condomínio: Millennium Residence
- CNPJ: 47.688.457/0001-02
- Endereço: Avenida Engenheiro Valdir Pedro Monachesi, 1.400, Aeroporto, Juiz de Fora/MG
- Presidente: {info_assembleia.get('presidente_nome', 'N/A')}
- Secretário: {info_assembleia.get('secretario_nome', 'N/A')}

NÃO inclua cabeçalhos, títulos principais ou assinaturas - apenas o conteúdo narrativo do trecho.

Trecho da transcrição:
{bloco_texto}

Transforme em texto formal seguindo o padrão:
"""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Você reescreve transcrições seguindo o padrão formal de atas da Contato Administração de Condomínios."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=4000,
        )
        return response.choices[0].message.content
    except Exception as e:
        raise RuntimeError(f"Erro ao gerar conteúdo formal: {e}")

def criar_paragrafo_encerramento(doc, info):
    """Cria o parágrafo de encerramento padrão"""
    doc.add_paragraph("\n")
    
    encerramento = doc.add_paragraph()
    encerramento.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    resultado_votacao = info.get('votacao_resultado', {})
    favoráveis = resultado_votacao.get('favoráveis', 'N/A')
    contrários = resultado_votacao.get('contrários', 'N/A') 
    abstenções = resultado_votacao.get('abstenções', 'N/A')
    
    if favoráveis != 'N/A':
        texto_votacao = f"com o seguinte resultado: {favoráveis} votos favoráveis, {contrários} votos contrários e {abstenções} abstenções. Alcançada a maioria simples, a proposta foi aprovada. "
    else:
        texto_votacao = ""
        
    data_encerramento = info['data_assembleia']
    
    texto_encerramento = f"""Nada mais havendo a tratar, encerrou-se a assembleia {texto_votacao}e lavrou-se a presente ata que vai assinada pela presidente da mesa e pelo secretário, para que se produzam seus efeitos legais e jurídicos. Juiz de Fora, {data_encerramento}."""
    
    encerramento.add_run(texto_encerramento)

def criar_assinaturas(doc, info):
    """Cria a seção de assinaturas"""
    doc.add_paragraph("\n\n")
    
    assinaturas = doc.add_paragraph()
    assinaturas.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = assinaturas.add_run(f"Presidente: {info['presidente_nome']}")
    run.font.bold = True
    
    assinaturas.add_run("\t\t\t")
    
    run = assinaturas.add_run(f"Secretário: {info['secretario_nome']}")
    run.font.bold = True

def gerar_ata_formal(transcricao, caminho_saida="ata_gerada.docx", status_callback=None, info_assembleia=None):
    """Função principal que gera a ata completa"""
    
    def report(msg):
        if status_callback:
            status_callback(msg)
        else:
            print(msg)
    
    try:
        # Usa informações fornecidas pelo usuário ou valores padrão
        if info_assembleia is None:
            report("Extraindo informações da assembleia...")
            info_assembleia = extrair_info_assembleia(transcricao)
        else:
            report("Usando informações fornecidas pelo usuário...")
        
        report("Criando documento...")
        doc = docx.Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Criar estrutura do documento
        report("Criando cabeçalho...")
        criar_cabecalho_documento(doc, info_assembleia)
        
        report("Criando título...")
        criar_titulo_ata(doc, info_assembleia['data_assembleia'], info_assembleia['tipo_assembleia'])
        
        report("Criando abertura...")
        criar_paragrafo_abertura(doc, info_assembleia)
        
        # Processar conteúdo em blocos
        report("Dividindo transcrição em blocos...")
        blocos = dividir_texto_em_blocos(transcricao, max_tokens=3500)
        report(f"Processando {len(blocos)} blocos de conteúdo...")
        
        for i, bloco in enumerate(blocos):
            report(f"Processando bloco {i+1}/{len(blocos)}...")
            conteudo_formal = gerar_conteudo_formal(bloco, info_assembleia)
            
            # Adicionar como parágrafo justificado
            paragrafo = doc.add_paragraph()
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragrafo.add_run(conteudo_formal)
        
        report("Criando encerramento...")
        criar_paragrafo_encerramento(doc, info_assembleia)
        
        report("Adicionando assinaturas...")
        criar_assinaturas(doc, info_assembleia)
        
        report("Salvando documento...")
        doc.save(caminho_saida)
        
        report(f"Ata gerada com sucesso: {caminho_saida}")
        
    except Exception as e:
        report(f"Erro ao gerar ata: {e}")
        raise